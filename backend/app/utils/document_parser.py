import os
import fitz  # PyMuPDF
import docx
import pdfplumber
from typing import Dict, Any, Optional
from fastapi import UploadFile, HTTPException
import re

class DocumentParser:
    """Parser for different document formats (PDF, DOCX, TXT)"""
    
    @staticmethod
    async def parse_document(file: UploadFile) -> Dict[str, Any]:
        """Parse uploaded document and extract content"""
        
        # Validate file type
        file_extension = os.path.splitext(file.filename)[1].lower()
        allowed_extensions = ['.pdf', '.docx', '.txt']
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type. Allowed: {', '.join(allowed_extensions)}"
            )
        
        try:
            if file_extension == '.pdf':
                return await DocumentParser._parse_pdf(file)
            elif file_extension == '.docx':
                return await DocumentParser._parse_docx(file)
            elif file_extension == '.txt':
                return await DocumentParser._parse_txt(file)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Error parsing document: {str(e)}")
    
    @staticmethod
    async def _parse_pdf(file: UploadFile) -> Dict[str, Any]:
        """Parse PDF document using PyMuPDF and pdfplumber for better text extraction"""
        
        # Read file content
        content = await file.read()
        
        # Try PyMuPDF first for better formatting preservation
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_content = ""
            
            for page in doc:
                text_content += page.get_text()
            
            doc.close()
            
            # If PyMuPDF doesn't extract much text, try pdfplumber
            if len(text_content.strip()) < 100:
                pdfplumber_doc = pdfplumber.open(content)
                text_content = ""
                for page in pdfplumber_doc.pages:
                    if page.extract_text():
                        text_content += page.extract_text() + "\n"
                pdfplumber_doc.close()
            
        except Exception:
            # Fallback to pdfplumber
            pdfplumber_doc = pdfplumber.open(content)
            text_content = ""
            for page in pdfplumber_doc.pages:
                if page.extract_text():
                    text_content += page.extract_text() + "\n"
            pdfplumber_doc.close()
        
        return {
            "content": text_content.strip(),
            "file_type": "pdf",
            "pages": len(fitz.open(stream=content, filetype="pdf"))
        }
    
    @staticmethod
    async def _parse_docx(file: UploadFile) -> Dict[str, Any]:
        """Parse DOCX document"""
        
        content = await file.read()
        
        # Save temporarily to parse with python-docx
        temp_path = f"/tmp/{file.filename}"
        with open(temp_path, "wb") as f:
            f.write(content)
        
        try:
            doc = docx.Document(temp_path)
            text_content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            os.remove(temp_path)
            
            return {
                "content": text_content.strip(),
                "file_type": "docx",
                "pages": DocumentParser._estimate_pages(text_content)
            }
            
        except Exception as e:
            if os.path.exists(temp_path):
                os.remove(temp_path)
            raise e
    
    @staticmethod
    async def _parse_txt(file: UploadFile) -> Dict[str, Any]:
        """Parse plain text document"""
        
        content = await file.read()
        text_content = content.decode('utf-8')
        
        return {
            "content": text_content.strip(),
            "file_type": "txt",
            "pages": DocumentParser._estimate_pages(text_content)
        }
    
    @staticmethod
    def _estimate_pages(text: str) -> int:
        """Estimate number of pages based on text length"""
        # Rough estimation: ~500 words per page
        words = len(text.split())
        return max(1, words // 500)
    
    @staticmethod
    def extract_resume_sections(text: str) -> Dict[str, str]:
        """Extract different sections from resume text"""
        
        sections = {
            "contact": "",
            "summary": "",
            "experience": "",
            "education": "",
            "skills": "",
            "projects": "",
            "certifications": ""
        }
        
        # Common section headers
        section_patterns = {
            "contact": r"(?i)(contact|personal|info|information)",
            "summary": r"(?i)(summary|objective|profile|about)",
            "experience": r"(?i)(experience|work|employment|career)",
            "education": r"(?i)(education|academic|degree|university|college)",
            "skills": r"(?i)(skills|technologies|tools|languages)",
            "projects": r"(?i)(projects|portfolio|achievements)",
            "certifications": r"(?i)(certifications|certificates|licenses)"
        }
        
        lines = text.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if line is a section header
            for section, pattern in section_patterns.items():
                if re.search(pattern, line) and len(line.split()) <= 3:
                    current_section = section
                    break
            
            # Add content to current section
            if current_section and line:
                sections[current_section] += line + "\n"
        
        # Clean up sections
        for section in sections:
            sections[section] = sections[section].strip()
        
        return sections 